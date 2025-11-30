"""
Generador Sección 5: Laboratorio
Tipo: 🟩 EXTRACCIÓN (datos desde MongoDB)
"""
from io import BytesIO
from pathlib import Path
from typing import Dict, Any, List, Optional
from .base import GeneradorSeccion
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate
import config
import logging

logger = logging.getLogger(__name__)


class GeneradorSeccion5(GeneradorSeccion):
    """Genera la sección 5: Laboratorio"""
    
    @property
    def nombre_seccion(self) -> str:
        return "5. LABORATORIO"
    
    @property
    def template_file(self) -> str:
        return "seccion_5_laboratorio.docx"
    
    def __init__(self, anio: int, mes: int, cargar_desde_mongodb: bool = True):
        super().__init__(anio, mes)
        self.cargar_desde_mongodb = cargar_desde_mongodb
        self.datos_laboratorio_raw: List[Dict] = []
        self.reporte_laboratorio: List[Dict] = []
        self.reintegrados_inventario: List[Dict] = []
        self.no_operativos: List[Dict] = []
        self.estado_garantia: List[Dict] = []
    
    def cargar_datos(self) -> None:
        """Carga datos de laboratorio desde MongoDB"""
        # Los datos se cargan desde MongoDB de forma asíncrona
        # Este método se llama desde el servicio después de cargar los datos
        pass
    
    def _calcular_cantidades_por_estado(self) -> List[Dict[str, Any]]:
        """
        Calcula las cantidades según los estados de los registros de laboratorio
        
        Returns:
            Lista de diccionarios con el reporte de laboratorio:
            [
                {"reporte": "REINTEGRADOS AL INVENTARIO", "cantidad": 10},
                {"reporte": "NO OPERATIVOS", "cantidad": 5},
                ...
            ]
        """
        if not self.datos_laboratorio_raw:
            logger.warning("No hay datos de laboratorio para calcular cantidades")
            return []
        
        # Estados para REINTEGRADOS AL INVENTARIO
        estados_reintegrados = ["OPERATIVO", "REPARADO", "REPARADO CAMPO"]
        
        # Contar registros por estado
        cantidad_reintegrados = 0
        cantidad_no_operativos = 0
        cantidad_garantia = 0
        cantidad_pendiente_parte = 0
        
        for registro in self.datos_laboratorio_raw:
            estado = str(registro.get("estado", "")).strip().upper()
            
            # REINTEGRADOS AL INVENTARIO
            if estado in estados_reintegrados:
                cantidad_reintegrados += 1
            
            # NO OPERATIVOS
            elif estado == "IRREPARABLE":
                cantidad_no_operativos += 1
            
            # ESTADO DE GARANTÍA
            elif estado == "ESTADO DE GARANTÍA":
                cantidad_garantia += 1
            
            # PENDIENTE POR PARTE
            elif estado == "PENDIENTE POR PARTE":
                cantidad_pendiente_parte += 1
        
        # Calcular total
        total = cantidad_reintegrados + cantidad_no_operativos + cantidad_garantia + cantidad_pendiente_parte
        
        # Construir lista de reporte
        reporte = [
            {
                "reporte": "REINTEGRADOS AL INVENTARIO",
                "cantidad": cantidad_reintegrados
            },
            {
                "reporte": "NO OPERATIVOS",
                "cantidad": cantidad_no_operativos
            },
            {
                "reporte": "ESTADO DE GARANTÍA",
                "cantidad": cantidad_garantia
            },
            {
                "reporte": "PENDIENTE POR PARTE",
                "cantidad": cantidad_pendiente_parte
            },
            {
                "reporte": "TOTAL",
                "cantidad": total
            }
        ]
        
        logger.info(f"Cantidades calculadas: Reintegrados={cantidad_reintegrados}, "
                   f"No Operativos={cantidad_no_operativos}, "
                   f"Garantía={cantidad_garantia}, "
                   f"Pendiente={cantidad_pendiente_parte}, "
                   f"Total={total}")
        
        return reporte
    
    def _filtrar_reintegrados_inventario(self) -> List[Dict[str, Any]]:
        """
        Filtra los registros que corresponden a REINTEGRADOS AL INVENTARIO
        
        Returns:
            Lista de registros con ESTADO = "OPERATIVO", "REPARADO" o "REPARADO CAMPO"
        """
        if not self.datos_laboratorio_raw:
            logger.warning("No hay datos de laboratorio para filtrar")
            return []
        
        # Estados para REINTEGRADOS AL INVENTARIO
        estados_reintegrados = ["OPERATIVO", "REPARADO", "REPARADO CAMPO"]
        
        reintegrados = []
        for registro in self.datos_laboratorio_raw:
            estado = str(registro.get("estado", "")).strip().upper()
            if estado in estados_reintegrados:
                reintegrados.append(registro)
        
        logger.info(f"Registros REINTEGRADOS AL INVENTARIO encontrados: {len(reintegrados)}")
        return reintegrados
    
    def _filtrar_no_operativos(self) -> List[Dict[str, Any]]:
        """
        Filtra los registros que corresponden a NO OPERATIVOS
        
        Returns:
            Lista de registros con ESTADO = "IRREPARABLE"
        """
        if not self.datos_laboratorio_raw:
            logger.warning("No hay datos de laboratorio para filtrar")
            return []
        
        # Estado para NO OPERATIVOS
        estado_no_operativo = "IRREPARABLE"
        
        no_operativos = []
        for registro in self.datos_laboratorio_raw:
            estado = str(registro.get("estado", "")).strip().upper()
            if estado == estado_no_operativo:
                no_operativos.append(registro)
        
        logger.info(f"Registros NO OPERATIVOS encontrados: {len(no_operativos)}")
        return no_operativos
    
    def _filtrar_estado_garantia(self) -> List[Dict[str, Any]]:
        """
        Filtra los registros que corresponden a ESTADO DE GARANTÍA
        
        Returns:
            Lista de registros con ESTADO = "ESTADO DE GARANTÍA"
        """
        if not self.datos_laboratorio_raw:
            logger.warning("No hay datos de laboratorio para filtrar")
            return []
        
        # Estado para GARANTÍA
        estado_garantia = "ESTADO DE GARANTÍA"
        
        garantia = []
        for registro in self.datos_laboratorio_raw:
            estado = str(registro.get("estado", "")).strip().upper()
            if estado == estado_garantia:
                garantia.append(registro)
        
        logger.info(f"Registros ESTADO DE GARANTÍA encontrados: {len(garantia)}")
        return garantia
    
    def _generar_texto_sin_garantia(self) -> str:
        """
        Genera el texto que se muestra cuando no hay registros de garantía
        
        Returns:
            Texto formateado con el mes y año
        """
        nombre_mes = config.MESES.get(self.mes, "MES").upper()
        return f"Bajo el trámite de RMA (Return Merchandise Authorization) para el periodo de corte comprendido del mes {nombre_mes} DE {self.anio}, no se tramita equipos bajo el proceso de garantía"
    
    def procesar(self) -> Dict[str, Any]:
        """Procesa los datos y retorna el contexto para el template"""
        # Calcular cantidades por estado
        self.reporte_laboratorio = self._calcular_cantidades_por_estado()
        
        # Filtrar registros REINTEGRADOS AL INVENTARIO
        self.reintegrados_inventario = self._filtrar_reintegrados_inventario()
        
        # Filtrar registros NO OPERATIVOS
        self.no_operativos = self._filtrar_no_operativos()
        
        # Filtrar registros ESTADO DE GARANTÍA
        self.estado_garantia = self._filtrar_estado_garantia()
        
        # Agregar marcadores de tabla al contexto
        contexto = {
            "TABLA_MARKER_REPORTE_LABORATORIO": "[[TABLA_REPORTE_LABORATORIO]]",
            "TABLA_MARKER_CONCEPTO_TECNICO": "[[TABLA_CONCEPTO_TECNICO]]",
            "TABLA_MARKER_CONCEPTO_TECNICO_NO_OPERATIVO": "[[TABLA_CONCEPTO_TECNICO_NO_OPERATIVO]]",
            "TABLA_MARKER_GARANTIA": "[[TABLA_GARANTIA]]",  # Marcador para la tabla de garantía
            "reporte_laboratorio": self.reporte_laboratorio,
            "reintegrados_inventario": self.reintegrados_inventario,
            "no_operativos": self.no_operativos,
            "estado_garantia": self.estado_garantia,
            "hay_garantia": len(self.estado_garantia) > 0,  # Variable booleana para condicionales
            "texto_sin_garantia": self._generar_texto_sin_garantia()
        }
        
        logger.info(f"Contexto generado: hay_garantia={contexto['hay_garantia']}, registros_garantia={len(self.estado_garantia)}")
        
        return contexto
    
    def generar(self):
        """
        Genera la sección completa con tablas dinámicas
        """
        # Cargar contexto base
        self.contexto = self.cargar_contexto_base()
        
        # Cargar datos específicos
        self.cargar_datos()
        
        # Procesar y agregar al contexto
        datos_seccion = self.procesar()
        self.contexto.update(datos_seccion)
        
        # Renderizar template con Jinja2
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template no encontrado: {self.template_path}")
        
        doc_template = DocxTemplate(self.template_path)
        doc_template.render(self.contexto)
        
        # Guardar el DocxTemplate a BytesIO para evitar problemas de XML
        buffer = BytesIO()
        doc_template.save(buffer)
        buffer.seek(0)
        
        # Abrir como Document de python-docx para manipulación de tablas
        doc = Document(buffer)
        
        # Reemplazar tabla de reporte de laboratorio
        self._reemplazar_tabla_por_marcador(
            doc,
            "TABLA_REPORTE_LABORATORIO",
            self.reporte_laboratorio,
            self._crear_tabla_reporte_laboratorio
        )
        
        # Reemplazar tabla de concepto técnico (REINTEGRADOS AL INVENTARIO)
        self._reemplazar_tabla_por_marcador(
            doc,
            "TABLA_CONCEPTO_TECNICO",
            self.reintegrados_inventario,
            self._crear_tabla_concepto_tecnico
        )
        
        # Reemplazar tabla de concepto técnico NO OPERATIVO
        self._reemplazar_tabla_por_marcador(
            doc,
            "TABLA_CONCEPTO_TECNICO_NO_OPERATIVO",
            self.no_operativos,
            self._crear_tabla_concepto_tecnico_no_operativo
        )
        
        # Reemplazar tabla/sección de GARANTÍA (con lógica condicional)
        self._reemplazar_tabla_o_texto_por_marcador(
            doc,
            "TABLA_GARANTIA",
            self.estado_garantia,
            self._crear_tabla_garantia,
            self._generar_texto_sin_garantia() if not self.estado_garantia else None
        )
        
        return doc
    
    def _reemplazar_tabla_por_marcador(self, doc: Document, marcador: str, datos: list, metodo_creacion) -> None:
        """
        Busca una tabla en el documento usando un marcador único y la reemplaza con datos.
        
        Args:
            doc: Documento de python-docx
            marcador: Nombre del marcador (ej: "TABLA_REPORTE_LABORATORIO")
            datos: Lista de datos para llenar la tabla
            metodo_creacion: Método que crea/llena la tabla
        """
        # El marcador se renderiza como [[TABLA_XXX]] después de procesar Jinja2
        marcador_renderizado = f"[[{marcador}]]"
        
        # Variaciones del marcador que pueden aparecer
        marcador_variaciones = [
            marcador_renderizado,  # Después de procesar Jinja2: [[TABLA_XXX]]
            f"{{{{ TABLA_MARKER_{marcador.replace('TABLA_', '')} }}}}",  # En el template original
            f"TABLA_MARKER_{marcador.replace('TABLA_', '')}",  # Variable de Jinja2 sin renderizar
            marcador.upper(),  # Nombre del marcador en mayúsculas
        ]
        
        # Normalizar el marcador para búsqueda
        marcador_busqueda = marcador.upper().replace("TABLA_", "")
        logger.info(f"Buscando tabla con marcador: {marcador_renderizado} (variaciones: {marcador_variaciones})")
        
        tabla_encontrada = None
        tabla_idx = None
        celda_con_marcador = None
        
        # Buscar el marcador en todas las tablas del documento
        for idx, tabla in enumerate(doc.tables):
            # Buscar en todas las celdas de la tabla
            for fila_idx, fila in enumerate(tabla.rows):
                for celda_idx, celda in enumerate(fila.cells):
                    # Obtener todo el texto de la celda (incluyendo párrafos)
                    texto_celda = ""
                    for parrafo in celda.paragraphs:
                        texto_celda += parrafo.text
                    
                    texto_celda_upper = texto_celda.upper()
                    
                    # Verificar si contiene el marcador renderizado [[TABLA_XXX]]
                    encontro_marcador = marcador_renderizado.upper() in texto_celda_upper
                    
                    # Si no se encuentra, buscar por variaciones
                    if not encontro_marcador:
                        for variacion in marcador_variaciones:
                            if variacion.upper() in texto_celda_upper:
                                encontro_marcador = True
                                break
                    
                    # También buscar por el nombre del marcador sin prefijos
                    if not encontro_marcador and marcador_busqueda in texto_celda_upper:
                        encontro_marcador = True
                    
                    if encontro_marcador:
                        tabla_encontrada = tabla
                        tabla_idx = idx
                        celda_con_marcador = (fila_idx, celda_idx)
                        logger.info(f"Marcador encontrado en tabla {idx}, fila {fila_idx}, celda {celda_idx}")
                        logger.debug(f"Texto de la celda: '{texto_celda[:100]}...'")
                        break
                
                if tabla_encontrada:
                    break
            
            if tabla_encontrada:
                break
        
        if tabla_encontrada and tabla_idx is not None:
            # Si no hay datos, no procesar la tabla (debe manejarse en otro método)
            if not datos:
                logger.warning(f"Tabla '{marcador}' encontrada pero no hay datos. Debe manejarse con _reemplazar_tabla_o_texto_por_marcador")
                return
            
            # Limpiar el marcador si está en el encabezado (fila 0)
            if celda_con_marcador:
                fila_idx, celda_idx = celda_con_marcador
                if fila_idx == 0:
                    # Si el marcador está en el encabezado, limpiarlo
                    celda = tabla_encontrada.rows[0].cells[celda_idx]
                    for parrafo in celda.paragraphs:
                        parrafo.clear()
            
            # Llamar al método de creación para reemplazar la tabla
            logger.info(f"Reemplazando tabla {tabla_idx} con {len(datos)} elementos")
            try:
                metodo_creacion(doc, tabla_encontrada)
                logger.info(f"Tabla '{marcador}' procesada correctamente")
            except Exception as e:
                logger.error(f"Error al procesar tabla '{marcador}': {str(e)}")
                import traceback
                traceback.print_exc()
        else:
            logger.warning(f"No se encontró tabla con marcador '{marcador}'")
            logger.info(f"Asegúrate de agregar '{{{{ TABLA_MARKER_{marcador.replace('TABLA_', '')} }}}}' en la primera celda de datos de la tabla en el template")
            logger.debug(f"Marcador buscado: {marcador_variaciones}")
    
    def _reemplazar_tabla_o_texto_por_marcador(
        self, 
        doc: Document, 
        marcador: str, 
        datos: list, 
        metodo_creacion_tabla,
        texto_sin_datos: str = None
    ) -> None:
        """
        Busca una tabla o área en el documento usando un marcador y:
        - Si hay datos: reemplaza la tabla con los datos
        - Si NO hay datos: reemplaza la tabla/área con un texto indicativo
        
        Args:
            doc: Documento de python-docx
            marcador: Nombre del marcador (ej: "TABLA_GARANTIA")
            datos: Lista de datos para llenar la tabla
            metodo_creacion_tabla: Método que crea/llena la tabla (solo si hay datos)
            texto_sin_datos: Texto a mostrar si no hay datos (opcional)
        """
        # Verificar si hay datos
        tiene_datos = datos and len(datos) > 0
        
        # Si no hay datos y hay texto sin datos, reemplazar con texto
        if not tiene_datos and texto_sin_datos:
            logger.info(f"No hay datos para '{marcador}' ({len(datos) if datos else 0} registros), reemplazando con texto indicativo")
            self._reemplazar_con_texto(doc, marcador, texto_sin_datos)
            return
        
        # Si hay datos, usar el método normal de reemplazar tabla
        if tiene_datos:
            logger.info(f"Hay {len(datos)} registros para '{marcador}', llenando tabla")
            self._reemplazar_tabla_por_marcador(doc, marcador, datos, metodo_creacion_tabla)
        else:
            logger.warning(f"No hay datos para '{marcador}' y no se proporcionó texto alternativo")
    
    def _reemplazar_con_texto(self, doc: Document, marcador: str, texto: str) -> None:
        """
        Reemplaza una tabla o área marcada con un texto simple
        
        Args:
            doc: Documento de python-docx
            marcador: Nombre del marcador
            texto: Texto a insertar
        """
        marcador_renderizado = f"[[{marcador}]]"
        marcador_variaciones = [
            marcador_renderizado,
            f"{{{{ TABLA_MARKER_{marcador.replace('TABLA_', '')} }}}}",
            f"TABLA_MARKER_{marcador.replace('TABLA_', '')}",
            marcador.upper(),
        ]
        
        marcador_busqueda = marcador.upper().replace("TABLA_", "")
        logger.info(f"Buscando área con marcador: {marcador_renderizado} para reemplazar con texto")
        
        # Buscar en tablas primero
        tabla_encontrada = None
        celda_con_marcador = None
        
        for idx, tabla in enumerate(doc.tables):
            for fila_idx, fila in enumerate(tabla.rows):
                for celda_idx, celda in enumerate(fila.cells):
                    texto_celda = ""
                    for parrafo in celda.paragraphs:
                        texto_celda += parrafo.text
                    
                    texto_celda_upper = texto_celda.upper()
                    encontro_marcador = marcador_renderizado.upper() in texto_celda_upper
                    
                    if not encontro_marcador:
                        for variacion in marcador_variaciones:
                            if variacion.upper() in texto_celda_upper:
                                encontro_marcador = True
                                break
                    
                    if not encontro_marcador and marcador_busqueda in texto_celda_upper:
                        encontro_marcador = True
                    
                    if encontro_marcador:
                        tabla_encontrada = tabla
                        celda_con_marcador = (fila_idx, celda_idx)
                        break
                
                if tabla_encontrada:
                    break
            
            if tabla_encontrada:
                break
        
        # Si encontramos una tabla, eliminar la tabla y reemplazarla con un párrafo
        if tabla_encontrada and celda_con_marcador:
            fila_idx, celda_idx = celda_con_marcador
            logger.info(f"Marcador encontrado en tabla {fila_idx}, celda {celda_idx}, eliminando tabla y reemplazando con párrafo")
            
            # Obtener el elemento XML de la tabla
            tbl_element = tabla_encontrada._tbl
            
            # Obtener el elemento body del documento
            body = doc.element.body
            
            # Buscar el índice de la tabla en el body
            tabla_idx = None
            for idx, elemento in enumerate(body):
                if elemento == tbl_element:
                    tabla_idx = idx
                    break
            
            if tabla_idx is not None:
                # Eliminar la tabla del documento
                body.remove(tbl_element)
                logger.info(f"Tabla eliminada del documento en posición {tabla_idx}")
                
                # Crear un nuevo párrafo usando el método del documento
                # Primero crear el párrafo temporalmente
                parrafo_temp = doc.add_paragraph()
                
                # Obtener el elemento XML del párrafo
                nuevo_p = parrafo_temp._p
                
                # Eliminar el párrafo temporal de su posición actual
                body.remove(nuevo_p)
                
                # Insertar el párrafo en la misma posición donde estaba la tabla
                body.insert(tabla_idx, nuevo_p)
                
                # Ahora usar el objeto Paragraph que ya creamos
                parrafo_obj = parrafo_temp
                
                # Agregar el texto con formato
                run = parrafo_obj.add_run(texto)
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                parrafo_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                logger.info(f"Párrafo creado en lugar de la tabla: '{texto[:50]}...'")
            else:
                # Si no encontramos el índice, intentar eliminar de otra forma
                logger.warning("No se encontró índice de tabla, usando método alternativo")
                try:
                    # Intentar eliminar directamente
                    parent = tbl_element.getparent()
                    if parent is not None:
                        parent.remove(tbl_element)
                        logger.info("Tabla eliminada usando getparent()")
                    
                    # Crear un nuevo párrafo al final
                    parrafo = doc.add_paragraph(texto)
                    run = parrafo.runs[0]
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    logger.info(f"Párrafo agregado: '{texto[:50]}...'")
                except Exception as e:
                    logger.error(f"Error al eliminar tabla: {str(e)}")
                    # Fallback: limpiar la tabla y dejar el texto en una celda fusionada
                    primera_fila = tabla_encontrada.rows[0]
                    primera_celda = primera_fila.cells[0]
                    
                    # Limpiar todas las celdas
                    for celda in primera_fila.cells:
                        for parrafo in celda.paragraphs:
                            parrafo.clear()
                    
                    # Eliminar filas adicionales
                    while len(tabla_encontrada.rows) > 1:
                        tbl = tabla_encontrada._tbl
                        tbl.remove(tabla_encontrada.rows[-1]._tr)
                    
                    # Fusionar columnas
                    if len(primera_fila.cells) > 1:
                        for i in range(len(primera_fila.cells) - 1, 0, -1):
                            try:
                                primera_celda._tc.merge(primera_fila.cells[i]._tc)
                            except:
                                pass
                    
                    # Agregar texto
                    parrafo = primera_celda.paragraphs[0] if primera_celda.paragraphs else primera_celda.add_paragraph()
                    parrafo.clear()
                    run = parrafo.add_run(texto)
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            return
        
        # Si no encontramos tabla, buscar en párrafos
        logger.info("No se encontró tabla, buscando en párrafos...")
        for parrafo in doc.paragraphs:
            texto_parrafo = parrafo.text
            texto_parrafo_upper = texto_parrafo.upper()
            
            encontro_marcador = marcador_renderizado.upper() in texto_parrafo_upper
            if not encontro_marcador:
                for variacion in marcador_variaciones:
                    if variacion.upper() in texto_parrafo_upper:
                        encontro_marcador = True
                        break
            
            if encontro_marcador or marcador_busqueda in texto_parrafo_upper:
                # Reemplazar el contenido del párrafo
                parrafo.clear()
                run = parrafo.add_run(texto)
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                logger.info(f"Párrafo reemplazado con texto: '{texto[:50]}...'")
                return
        
        logger.warning(f"No se encontró marcador '{marcador}' para reemplazar con texto")
    
    def _crear_tabla_reporte_laboratorio(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos del reporte de laboratorio
        
        Args:
            doc: Documento de python-docx
            tabla_existente: Tabla existente en el documento
        """
        logger.info(f"Reemplazando tabla de reporte de laboratorio con {len(self.reporte_laboratorio)} filas")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        logger.info(f"Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener número de columnas
        num_cols = len(tabla_existente.columns)
        logger.info(f"Tabla tiene {num_cols} columnas")
        
        # Verificar/actualizar encabezados
        if len(tabla_existente.rows) > 0:
            encabezados_esperados = ["REPORTE LABORATORIO", "CANTIDAD"]
            primera_fila = tabla_existente.rows[0]
            
            # Actualizar encabezados si no coinciden
            for i, header_text in enumerate(encabezados_esperados):
                if i < num_cols and i < len(primera_fila.cells):
                    celda = primera_fila.cells[i]
                    texto_actual = celda.text.strip().upper()
                    if texto_actual != header_text.upper():
                        celda.text = header_text
                        # Formatear encabezado
                        for parrafo in celda.paragraphs:
                            for run in parrafo.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
                            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar filas con los datos
        for item in self.reporte_laboratorio:
            fila = tabla_existente.add_row()
            celdas = fila.cells
            
            # Columna 1: REPORTE LABORATORIO
            if len(celdas) > 0:
                celda_reporte = celdas[0]
                celda_reporte.text = item.get("reporte", "")
                # Formatear celda
                for parrafo in celda_reporte.paragraphs:
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in parrafo.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
            
            # Columna 2: CANTIDAD
            if len(celdas) > 1:
                celda_cantidad = celdas[1]
                cantidad = item.get("cantidad", 0)
                celda_cantidad.text = str(cantidad)
                # Formatear celda
                for parrafo in celda_cantidad.paragraphs:
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in parrafo.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
                        # Si es TOTAL, poner en negrita
                        if item.get("reporte", "").upper() == "TOTAL":
                            run.font.bold = True
        
        logger.info(f"Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.reporte_laboratorio)} datos)")
    
    def _crear_tabla_concepto_tecnico(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos de REINTEGRADOS AL INVENTARIO
        
        Args:
            doc: Documento de python-docx
            tabla_existente: Tabla existente en el documento
        """
        logger.info(f"Reemplazando tabla de concepto técnico con {len(self.reintegrados_inventario)} registros")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        logger.info(f"Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener número de columnas
        num_cols = len(tabla_existente.columns)
        logger.info(f"Tabla tiene {num_cols} columnas")
        
        # Verificar/actualizar encabezados
        if len(tabla_existente.rows) > 0:
            encabezados_esperados = ["ID", "FECHA", "PUNTO", "EQUIPO", "SERIAL", "ESTADO"]
            primera_fila = tabla_existente.rows[0]
            
            # Si la tabla tiene menos de 6 columnas, intentar agregar columnas faltantes
            if num_cols < 6:
                logger.warning(f"La tabla tiene solo {num_cols} columnas, se necesitan 6. Intentando agregar columnas...")
                from docx.shared import Inches
                columnas_faltantes = 6 - num_cols
                for i in range(columnas_faltantes):
                    tabla_existente.add_column(Inches(1.5))
                    logger.info(f"Columna {num_cols + i + 1} agregada")
                num_cols = len(tabla_existente.columns)
                logger.info(f"Tabla ahora tiene {num_cols} columnas")
            
            # Actualizar encabezados si no coinciden
            for i, header_text in enumerate(encabezados_esperados):
                if i < num_cols and i < len(primera_fila.cells):
                    celda = primera_fila.cells[i]
                    texto_actual = celda.text.strip().upper()
                    if texto_actual != header_text.upper():
                        celda.text = header_text
                        # Formatear encabezado
                        for parrafo in celda.paragraphs:
                            for run in parrafo.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
                            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Mapeo de campos de los registros a las columnas
        mapeo_columnas = {
            0: 'id',
            1: 'fecha',
            2: 'punto',
            3: 'equipo',
            4: 'serial',
            5: 'estado'
        }
        
        # Agregar filas con los datos
        for idx, registro in enumerate(self.reintegrados_inventario, start=1):
            fila = tabla_existente.add_row()
            celdas = fila.cells
            
            # Llenar cada celda según el mapeo
            for i in range(min(num_cols, 6)):
                campo = mapeo_columnas.get(i, '')
                
                # Para la columna ID (índice 0), usar el consecutivo en lugar del ID del registro
                if i == 0:
                    valor = str(idx)  # Consecutivo empezando desde 1
                else:
                    valor = registro.get(campo, '')
                
                if i < len(celdas):
                    # Limpiar contenido existente de la celda
                    celdas[i].text = ''
                    # Agregar el nuevo texto
                    parrafo = celdas[i].paragraphs[0] if celdas[i].paragraphs else celdas[i].add_paragraph()
                    parrafo.clear()
                    run = parrafo.add_run(str(valor) if valor else '')
                    
                    # Aplicar formato según la columna
                    if i == 0:  # ID
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 1:  # FECHA
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 2:  # PUNTO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 3:  # EQUIPO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 4:  # SERIAL
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 5:  # ESTADO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    
                    run.font.name = 'Calibri'
        
        logger.info(f"Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.reintegrados_inventario)} datos)")
    
    def _crear_tabla_concepto_tecnico_no_operativo(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos de NO OPERATIVOS (IRREPARABLE)
        
        Args:
            doc: Documento de python-docx
            tabla_existente: Tabla existente en el documento
        """
        logger.info(f"Reemplazando tabla de concepto técnico NO OPERATIVO con {len(self.no_operativos)} registros")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        logger.info(f"Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener número de columnas
        num_cols = len(tabla_existente.columns)
        logger.info(f"Tabla tiene {num_cols} columnas")
        
        # Verificar/actualizar encabezados
        if len(tabla_existente.rows) > 0:
            encabezados_esperados = ["ID", "FECHA", "PUNTO", "EQUIPO", "SERIAL", "ESTADO"]
            primera_fila = tabla_existente.rows[0]
            
            # Si la tabla tiene menos de 6 columnas, intentar agregar columnas faltantes
            if num_cols < 6:
                logger.warning(f"La tabla tiene solo {num_cols} columnas, se necesitan 6. Intentando agregar columnas...")
                from docx.shared import Inches
                columnas_faltantes = 6 - num_cols
                for i in range(columnas_faltantes):
                    tabla_existente.add_column(Inches(1.5))
                    logger.info(f"Columna {num_cols + i + 1} agregada")
                num_cols = len(tabla_existente.columns)
                logger.info(f"Tabla ahora tiene {num_cols} columnas")
            
            # Actualizar encabezados si no coinciden
            for i, header_text in enumerate(encabezados_esperados):
                if i < num_cols and i < len(primera_fila.cells):
                    celda = primera_fila.cells[i]
                    texto_actual = celda.text.strip().upper()
                    if texto_actual != header_text.upper():
                        celda.text = header_text
                        # Formatear encabezado
                        for parrafo in celda.paragraphs:
                            for run in parrafo.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
                            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Mapeo de campos de los registros a las columnas
        mapeo_columnas = {
            0: 'id',
            1: 'fecha',
            2: 'punto',
            3: 'equipo',
            4: 'serial',
            5: 'estado'
        }
        
        # Agregar filas con los datos
        for idx, registro in enumerate(self.no_operativos, start=1):
            fila = tabla_existente.add_row()
            celdas = fila.cells
            
            # Llenar cada celda según el mapeo
            for i in range(min(num_cols, 6)):
                campo = mapeo_columnas.get(i, '')
                
                # Para la columna ID (índice 0), usar el consecutivo en lugar del ID del registro
                if i == 0:
                    valor = str(idx)  # Consecutivo empezando desde 1
                else:
                    valor = registro.get(campo, '')
                
                if i < len(celdas):
                    # Limpiar contenido existente de la celda
                    celdas[i].text = ''
                    # Agregar el nuevo texto
                    parrafo = celdas[i].paragraphs[0] if celdas[i].paragraphs else celdas[i].add_paragraph()
                    parrafo.clear()
                    run = parrafo.add_run(str(valor) if valor else '')
                    
                    # Aplicar formato según la columna
                    if i == 0:  # ID
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 1:  # FECHA
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 2:  # PUNTO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 3:  # EQUIPO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 4:  # SERIAL
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 5:  # ESTADO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    
                    run.font.name = 'Calibri'
        
        logger.info(f"Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.no_operativos)} datos)")
    
    def _crear_tabla_garantia(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos de ESTADO DE GARANTÍA
        
        Args:
            doc: Documento de python-docx
            tabla_existente: Tabla existente en el documento
        """
        logger.info(f"Reemplazando tabla de garantía con {len(self.estado_garantia)} registros")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        logger.info(f"Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener número de columnas
        num_cols = len(tabla_existente.columns)
        logger.info(f"Tabla tiene {num_cols} columnas")
        
        # Verificar/actualizar encabezados
        if len(tabla_existente.rows) > 0:
            encabezados_esperados = ["ID", "FECHA", "PUNTO", "EQUIPO", "SERIAL", "ESTADO"]
            primera_fila = tabla_existente.rows[0]
            
            # Si la tabla tiene menos de 6 columnas, intentar agregar columnas faltantes
            if num_cols < 6:
                logger.warning(f"La tabla tiene solo {num_cols} columnas, se necesitan 6. Intentando agregar columnas...")
                from docx.shared import Inches
                columnas_faltantes = 6 - num_cols
                for i in range(columnas_faltantes):
                    tabla_existente.add_column(Inches(1.5))
                    logger.info(f"Columna {num_cols + i + 1} agregada")
                num_cols = len(tabla_existente.columns)
                logger.info(f"Tabla ahora tiene {num_cols} columnas")
            
            # Actualizar encabezados si no coinciden
            for i, header_text in enumerate(encabezados_esperados):
                if i < num_cols and i < len(primera_fila.cells):
                    celda = primera_fila.cells[i]
                    texto_actual = celda.text.strip().upper()
                    if texto_actual != header_text.upper():
                        celda.text = header_text
                        # Formatear encabezado
                        for parrafo in celda.paragraphs:
                            for run in parrafo.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
                            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Mapeo de campos de los registros a las columnas
        mapeo_columnas = {
            0: 'id',
            1: 'fecha',
            2: 'punto',
            3: 'equipo',
            4: 'serial',
            5: 'estado'
        }
        
        # Agregar filas con los datos
        for idx, registro in enumerate(self.estado_garantia, start=1):
            fila = tabla_existente.add_row()
            celdas = fila.cells
            
            # Llenar cada celda según el mapeo
            for i in range(min(num_cols, 6)):
                campo = mapeo_columnas.get(i, '')
                
                # Para la columna ID (índice 0), usar el consecutivo en lugar del ID del registro
                if i == 0:
                    valor = str(idx)  # Consecutivo empezando desde 1
                else:
                    valor = registro.get(campo, '')
                
                if i < len(celdas):
                    # Limpiar contenido existente de la celda
                    celdas[i].text = ''
                    # Agregar el nuevo texto
                    parrafo = celdas[i].paragraphs[0] if celdas[i].paragraphs else celdas[i].add_paragraph()
                    parrafo.clear()
                    run = parrafo.add_run(str(valor) if valor else '')
                    
                    # Aplicar formato según la columna
                    if i == 0:  # ID
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 1:  # FECHA
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 2:  # PUNTO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 3:  # EQUIPO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 4:  # SERIAL
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 5:  # ESTADO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    
                    run.font.name = 'Calibri'
        
        logger.info(f"Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.estado_garantia)} datos)")
    
    def guardar(self, output_path: Path) -> None:
        """
        Genera y guarda la sección, asegurando que los cambios en las tablas se guarden correctamente
        """
        doc = self.generar()
        
        # Guardar el documento
        doc.save(str(output_path))
        logger.info(f"{self.nombre_seccion} guardada en: {output_path}")
