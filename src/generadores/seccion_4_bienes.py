"""
Generador Sección 4: Informe de Bienes y Servicios
Tipo: 🟨 ANÁLISIS IA (párrafos) + 🟩 EXTRACCIÓN DATOS (Excel/SharePoint)

Subsecciones:
- 4.1 Gestión de Inventario
- 4.2 Entradas Almacén SDSCJ
- 4.3 Entrega Equipos No Operativos Almacén SDSCJ
- 4.4 Gestiones de Inclusión a la Bolsa

Características especiales:
- Conversión de valores a letras (ej: "DOSCIENTOS CUARENTA Y CINCO MIL PESOS M/CTE")
- Tablas de ítems con cantidades y valores
- Referencias a comunicados y anexos
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
import json
from .base import GeneradorSeccion
import config
from src.utils.formato_moneda import numero_a_letras, formato_moneda_cop
from src.extractores.excel_extractor import get_excel_extractor


class GeneradorSeccion4(GeneradorSeccion):
    """Genera la Sección 4: Informe de Bienes y Servicios"""
    
    # Colores
    COLOR_AZUL_OSCURO = RGBColor(31, 78, 121)
    COLOR_AZUL_MEDIO = RGBColor(46, 117, 182)
    COLOR_GRIS = RGBColor(64, 64, 64)
    COLOR_VERDE = RGBColor(0, 128, 0)
    
    @property
    def nombre_seccion(self) -> str:
        return "4. INFORME DE BIENES Y SERVICIOS"
    
    @property
    def template_file(self) -> str:
        return "seccion_4_bienes_servicios.docx"
    
    def __init__(self, anio: int, mes: int):
        super().__init__(anio, mes)
        self.datos: Dict[str, Any] = {}
        self.doc: Optional[Document] = None
        self.excel_extractor = get_excel_extractor()
    
    def _configurar_estilos(self):
        """Configura los estilos del documento"""
        if self.doc is None:
            return
        
        # Configurar márgenes (1.27 cm = 0.5 pulgadas)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
        
        # Estilo normal
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Título Sección (4.)
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = self.COLOR_AZUL_OSCURO
        
        # Subsecciones (4.1-4.4)
        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.font.color.rgb = self.COLOR_AZUL_MEDIO
        
        # Subtítulos
        h3 = self.doc.styles['Heading 3']
        h3.font.name = 'Arial'
        h3.font.size = Pt(11)
        h3.font.bold = True
        h3.font.color.rgb = self.COLOR_GRIS
    
    def _set_cell_shading(self, cell, color_hex: str):
        """Establece el color de fondo de una celda"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _centrar_celda_vertical(self, cell):
        """Centra verticalmente el contenido de una celda"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
    
    def _numero_a_letras(self, numero: float, incluir_moneda: bool = True) -> str:
        """
        Convierte un número a su representación en letras en español
        Usa la utilidad centralizada
        """
        return numero_a_letras(numero, incluir_moneda)
    
    def _formato_moneda(self, numero: float) -> str:
        """Formatea número como moneda colombiana: $245.678.910"""
        return formato_moneda_cop(numero)
    
    def _agregar_tabla(self, encabezados: list, filas: list, anchos: list = None,
                       alineaciones: list = None, fila_total: list = None):
        """Agrega una tabla con formato profesional"""
        tabla = self.doc.add_table(rows=1, cols=len(encabezados))
        tabla.style = 'Table Grid'
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados
        hdr_cells = tabla.rows[0].cells
        for i, texto in enumerate(encabezados):
            hdr_cells[i].text = texto
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            self._set_cell_shading(hdr_cells[i], "1F4E79")
            self._centrar_celda_vertical(hdr_cells[i])
        
        # Filas de datos
        for fila_datos in filas:
            row_cells = tabla.add_row().cells
            for i, texto in enumerate(fila_datos):
                row_cells[i].text = str(texto) if texto is not None else ""
                for paragraph in row_cells[i].paragraphs:
                    if alineaciones and i < len(alineaciones):
                        paragraph.alignment = alineaciones[i]
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                self._centrar_celda_vertical(row_cells[i])
        
        # Fila de total (si se proporciona)
        if fila_total:
            row_cells = tabla.add_row().cells
            for i, texto in enumerate(fila_total):
                row_cells[i].text = str(texto) if texto is not None else ""
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                self._set_cell_shading(row_cells[i], "D9E1F2")  # Azul claro
                self._centrar_celda_vertical(row_cells[i])
        
        # Ajustar anchos
        if anchos:
            for i, ancho in enumerate(anchos):
                for cell in tabla.columns[i].cells:
                    cell.width = Inches(ancho)
        
        self.doc.add_paragraph()
        return tabla
    
    def _obtener_dias_mes(self) -> str:
        """Obtiene el último día del mes"""
        from calendar import monthrange
        ultimo_dia = monthrange(self.anio, self.mes)[1]
        return f"{ultimo_dia:02d}"
    
    def _agregar_parrafo(self, texto: str, justificado: bool = True, 
                         negrita: bool = False, italica: bool = False):
        """Agrega un párrafo de texto"""
        p = self.doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = negrita
        run.italic = italica
        if justificado:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        return p
    
    def _seccion_4_1_gestion_inventario(self):
        """4.1 Gestión de Inventario"""
        self.doc.add_heading("4.1. GESTIÓN DE INVENTARIO", level=2)
        
        gestion = self.datos.get('gestion_inventario', {})
        texto = gestion.get('texto', '')
        ruta = gestion.get('ruta', '')
        
        # Si hay ruta, reemplazar {{root_mes}} o {{01NOV - 30NOV}} con el mes correspondiente
        if ruta:
            mes_nombre = config.MESES[self.mes]
            # Reemplazar variables de mes en la ruta
            ruta = ruta.replace('{{root_mes}}', f"01{mes_nombre.upper()[:3]} - 30{mes_nombre.upper()[:3]}")
            # Buscar patrones como {{01NOV - 30NOV}} y reemplazarlos
            import re
            patron = r'\{\{01\w+ - 30\w+\}\}'
            ruta = re.sub(patron, f"01{mes_nombre.upper()[:3]} - 30{mes_nombre.upper()[:3]}", ruta)
            
            # Combinar texto y ruta
            if texto:
                texto_completo = f"{texto}\n\n{ruta}"
            else:
                texto_completo = ruta
        else:
            texto_completo = texto
        
        if texto_completo:
            self._agregar_parrafo(texto_completo)
    
    def _seccion_4_2_entradas_almacen(self):
        """4.2 Entradas Almacén SDSCJ"""
        self.doc.add_heading("4.2. ENTRADAS ALMACÉN SDSCJ", level=2)
        
        mes = config.MESES[self.mes]
        anio = self.anio
        entradas = self.datos.get('entradas_almacen', {})
        
        # Datos del comunicado
        comunicado = entradas.get('comunicado', {})
        items = entradas.get('items', [])
        texto_entradas = entradas.get('texto', '')
        
        # Si hay texto personalizado, usarlo; si no, generar texto automático
        if texto_entradas:
            self._agregar_parrafo(texto_entradas)
        else:
            numero = comunicado.get('numero', 'N/A')
            titulo = comunicado.get('titulo', 'N/A')
            fecha = comunicado.get('fecha', 'N/A')
            if numero != 'N/A' or fecha != 'N/A':
                texto = f"""Se relacionan las entradas al almacén SDSCJ durante el mes correspondiente, donde se registran los equipos y componentes recibidos para el sistema de videovigilancia."""
                if numero != 'N/A':
                    texto += f" Comunicado: {numero}"
                if fecha != 'N/A':
                    texto += f" Fecha: {fecha}"
                texto += "."
                self._agregar_parrafo(texto)
        
        valor_total = sum(item.get('valor_total', 0) for item in items)
        
        # Tabla de ítems
        if items:
            self._agregar_parrafo("Elementos ingresados al almacén:", negrita=True)
            
            filas = []
            for idx, item in enumerate(items, 1):
                filas.append([
                    idx,
                    item.get('descripcion', ''),
                    item.get('cantidad', 0),
                    item.get('unidad', 'UN'),
                    self._formato_moneda(item.get('valor_unitario', 0)) if item.get('valor_unitario', 0) > 0 else "-",
                    self._formato_moneda(item.get('valor_total', 0)) if item.get('valor_total', 0) > 0 else "-"
                ])
            
            # Alineaciones: No., Descripción, Cant, Unidad, V.Unit, V.Total
            alineaciones = [
                WD_ALIGN_PARAGRAPH.CENTER,  # No.
                WD_ALIGN_PARAGRAPH.LEFT,    # Descripción
                WD_ALIGN_PARAGRAPH.CENTER,  # Cantidad
                WD_ALIGN_PARAGRAPH.CENTER,  # Unidad
                WD_ALIGN_PARAGRAPH.RIGHT,   # Valor Unit
                WD_ALIGN_PARAGRAPH.RIGHT,   # Valor Total
            ]
            
            # Fila de total solo si hay valores
            fila_total = None
            if valor_total > 0:
                fila_total = ["", "TOTAL", "", "", "", self._formato_moneda(valor_total)]
            
            self._agregar_tabla(
                encabezados=["No.", "DESCRIPCIÓN", "CANT.", "UND", "VALOR UNIT.", "VALOR TOTAL"],
                filas=filas,
                anchos=[0.4, 2.5, 0.5, 0.5, 1.0, 1.0],  # Según especificaciones del prompt
                alineaciones=alineaciones,
                fila_total=fila_total
            )
        
        # Valor en letras según template del prompt
        if valor_total > 0:
            valor_letras = self._numero_a_letras(valor_total)
            valor_numerico = self._formato_moneda(valor_total)
            self._agregar_parrafo(
                f"El valor total de los elementos ingresados asciende a {valor_letras} ({valor_numerico}).",
                negrita=True
            )
        
        # Anexos
        anexos = entradas.get('anexos', [])
        if anexos:
            self._agregar_parrafo("Se adjuntan como anexos:", negrita=True)
            for anexo in anexos:
                self._agregar_parrafo(f"• {anexo}")
    
    def _seccion_4_3_equipos_no_operativos(self):
        """4.3 Entrega Equipos No Operativos Almacén SDSCJ"""
        self.doc.add_heading("4.3. ENTREGA EQUIPOS NO OPERATIVOS ALMACÉN SDSCJ", level=2)
        
        mes = config.MESES[self.mes]
        anio = self.anio
        equipos_data = self.datos.get('equipos_no_operativos', {})
        
        comunicado = equipos_data.get('comunicado', {})
        equipos = equipos_data.get('equipos', [])
        texto_equipos = equipos_data.get('texto', '')
        valor_total = sum(eq.get('valor', 0) for eq in equipos)
        
        # Si hay texto personalizado, usarlo; si no, generar texto automático
        if texto_equipos:
            self._agregar_parrafo(texto_equipos)
        else:
            # Párrafo introductorio
            texto = f"""Se realiza el trámite de entrega de equipos no operativos al almacén SDSCJ en el mes de {mes} del {anio}"""
            if comunicado.get('numero'):
                texto += f" bajo el comunicado {comunicado.get('numero', 'N/A')}"
            if comunicado.get('fecha'):
                texto += f" el {comunicado.get('fecha', 'N/A')}"
            texto += ", para concepto técnico y disposición final."
            self._agregar_parrafo(texto)
        
        # Tabla de equipos
        if equipos:
            self._agregar_parrafo("Equipos entregados:", negrita=True)
            
            filas = []
            for idx, eq in enumerate(equipos, 1):
                filas.append([
                    idx,
                    eq.get('descripcion', ''),
                    eq.get('serial', '') if eq.get('serial') else '-',
                    eq.get('cantidad', 1),
                    eq.get('motivo', ''),
                    self._formato_moneda(eq.get('valor', 0)) if eq.get('valor', 0) > 0 else "-"
                ])
            
            alineaciones = [
                WD_ALIGN_PARAGRAPH.CENTER,  # No.
                WD_ALIGN_PARAGRAPH.LEFT,    # Descripción
                WD_ALIGN_PARAGRAPH.LEFT,    # Serial
                WD_ALIGN_PARAGRAPH.CENTER,  # Cantidad
                WD_ALIGN_PARAGRAPH.LEFT,    # Motivo
                WD_ALIGN_PARAGRAPH.RIGHT,   # Valor
            ]
            
            fila_total = None
            if valor_total > 0:
                fila_total = ["", "TOTAL", "", "", "", self._formato_moneda(valor_total)]
            
            self._agregar_tabla(
                encabezados=["No.", "DESCRIPCIÓN", "SERIAL", "CANT.", "MOTIVO", "VALOR"],
                filas=filas,
                anchos=[0.4, 2.0, 1.0, 0.5, 1.5, 1.0],  # Según especificaciones del prompt
                alineaciones=alineaciones,
                fila_total=fila_total
            )
        
        # Valor en letras (agregar según especificaciones)
        if valor_total > 0:
            valor_letras = self._numero_a_letras(valor_total)
            valor_numerico = self._formato_moneda(valor_total)
            self._agregar_parrafo(
                f"El valor total de los equipos entregados asciende a {valor_letras} ({valor_numerico}).",
                negrita=True
            )
        
        # Anexos
        anexos = equipos_data.get('anexos', [])
        if anexos:
            self._agregar_parrafo("Se adjuntan como anexos:", negrita=True)
            for anexo in anexos:
                self._agregar_parrafo(f"• {anexo}")
    
    def _seccion_4_4_inclusiones_bolsa(self):
        """4.4 Gestiones de Inclusión a la Bolsa"""
        self.doc.add_heading("4.4. GESTIONES DE INCLUSIÓN A LA BOLSA", level=2)
        
        mes = config.MESES[self.mes]
        anio = self.anio
        inclusiones_data = self.datos.get('inclusiones_bolsa', {})
        
        comunicado = inclusiones_data.get('comunicado', {})
        items = inclusiones_data.get('items', [])
        estado = inclusiones_data.get('estado', 'Sin solicitudes')
        texto_inclusiones = inclusiones_data.get('texto', '')
        valor_total = sum(item.get('valor_total', 0) for item in items)
        
        # Si hay texto personalizado, usarlo; si no, generar texto automático
        if texto_inclusiones:
            self._agregar_parrafo(texto_inclusiones)
        else:
            # Párrafo introductorio
            texto = f"""Para el periodo comprendido entre el 01 al {self._obtener_dias_mes()} de {mes} del {anio} se realizó la presentación de solicitud de inclusión de bolsa de repuestos de:"""
            self._agregar_parrafo(texto)
        
        # Estado de la solicitud solo si no es "Sin solicitudes"
        if estado != 'Sin solicitudes':
            self._agregar_parrafo(f"Estado de la solicitud: {estado}", negrita=True)
        
        # Tabla de ítems
        if items:
            self._agregar_parrafo("Elementos solicitados para inclusión:", negrita=True)
            
            filas = []
            for idx, item in enumerate(items, 1):
                filas.append([
                    idx,
                    item.get('descripcion', ''),
                    item.get('cantidad', 0),
                    item.get('unidad', 'UN'),
                    self._formato_moneda(item.get('valor_unitario', 0)),
                    self._formato_moneda(item.get('valor_total', 0)),
                    item.get('justificacion', '')
                ])
            
            alineaciones = [
                WD_ALIGN_PARAGRAPH.CENTER,  # No.
                WD_ALIGN_PARAGRAPH.LEFT,    # Descripción
                WD_ALIGN_PARAGRAPH.CENTER,  # Cantidad
                WD_ALIGN_PARAGRAPH.CENTER,  # Unidad
                WD_ALIGN_PARAGRAPH.RIGHT,   # Valor Unit
                WD_ALIGN_PARAGRAPH.RIGHT,   # Valor Total
                WD_ALIGN_PARAGRAPH.LEFT,    # Justificación
            ]
            
            fila_total = None
            if valor_total > 0:
                fila_total = ["", "TOTAL", "", "", "", self._formato_moneda(valor_total), ""]
            
            self._agregar_tabla(
                encabezados=["No.", "DESCRIPCIÓN", "CANT.", "UND", "VALOR UNIT.", "VALOR TOTAL", "JUSTIFICACIÓN"],
                filas=filas,
                anchos=[0.4, 2.0, 0.5, 0.5, 1.0, 1.0, 1.5],
                alineaciones=alineaciones,
                fila_total=fila_total
            )
        
        # Valor en letras
        if valor_total > 0:
            valor_letras = self._numero_a_letras(valor_total)
            self._agregar_parrafo(
                f"El valor total de la solicitud de inclusión asciende a {valor_letras} ({self._formato_moneda(valor_total)}).",
                negrita=True
            )
        
        # Anexos
        anexos = inclusiones_data.get('anexos', [])
        if anexos:
            self._agregar_parrafo("Se adjuntan como anexos:", negrita=True)
            for anexo in anexos:
                self._agregar_parrafo(f"• {anexo}")
    
    def cargar_datos(self) -> None:
        """
        Carga los datos específicos de la sección 4 desde JSON y Excel
        Si ya hay datos cargados (desde MongoDB), no los sobrescribe
        """
        # Si ya hay datos cargados (desde MongoDB), no cargar desde archivos
        if self.datos and any(key in self.datos for key in ['gestion_inventario', 'entradas_almacen', 'equipos_no_operativos', 'inclusiones_bolsa']):
            return
        
        # Cargar datos desde archivo JSON
        # Intentar primero con formato numérico, luego con nombre de mes
        archivo = config.FUENTES_DIR / f"bienes_{self.mes}_{self.anio}.json"
        if not archivo.exists():
            archivo = config.FUENTES_DIR / f"bienes_{config.MESES[self.mes].lower()}_{self.anio}.json"
        
        if archivo.exists():
            try:
                with open(archivo, 'r', encoding='utf-8') as f:
                    self.datos = json.load(f)
            except Exception as e:
                print(f"[WARNING] Error al cargar datos desde {archivo}: {e}")
                self.datos = {}
        else:
            print(f"[WARNING] Archivo de datos no encontrado: {archivo}")
            self.datos = {}
        
        # Cargar datos desde Excel (usando extractor)
        entradas_excel = self.excel_extractor.get_entradas_almacen(self.anio, self.mes)
        if entradas_excel and entradas_excel.get('items'):
            self.datos['entradas_almacen'] = entradas_excel
        
        equipos_excel = self.excel_extractor.get_equipos_no_operativos(self.anio, self.mes)
        if equipos_excel and equipos_excel.get('equipos'):
            self.datos['equipos_no_operativos'] = equipos_excel
        
        inclusiones_excel = self.excel_extractor.get_inclusiones_bolsa(self.anio, self.mes)
        if inclusiones_excel and inclusiones_excel.get('items'):
            self.datos['inclusiones_bolsa'] = inclusiones_excel
        
        # Agregar mes y año a los datos para compatibilidad
        self.datos['mes'] = config.MESES[self.mes]
        self.datos['anio'] = self.anio
    
    def procesar(self) -> Dict[str, Any]:
        """Procesa los datos y retorna el contexto para el template"""
        # Esta clase no usa templates, pero debe implementar el método abstracto
        return {}
    
    def generar(self) -> Document:
        """
        Genera el documento completo de la Sección 4
        Sobrescribe el método de la clase base para usar python-docx directamente
        """
        # Cargar datos si no se han cargado
        if not self.datos:
            self.cargar_datos()
        
        # Crear documento
        self.doc = Document()
        self._configurar_estilos()
        
        # Título principal
        self.doc.add_heading("4. INFORME DE BIENES Y SERVICIOS", level=1)
        
        # Generar subsecciones
        self._seccion_4_1_gestion_inventario()
        self._seccion_4_2_entradas_almacen()
        self._seccion_4_3_equipos_no_operativos()
        self._seccion_4_4_inclusiones_bolsa()
        
        # Separador fin de sección
        self.doc.add_paragraph()
        p = self.doc.add_paragraph("═" * 60)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = self.doc.add_paragraph("Fin Sección 4 - Informe de Bienes y Servicios")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        return self.doc
    
    def guardar(self, output_path: Path) -> None:
        """
        Genera y guarda la sección
        Sobrescribe el método de la clase base para usar python-docx
        """
        if self.doc is None:
            self.generar()
        
        self.doc.save(str(output_path))
        print(f"[OK] {self.nombre_seccion} guardada en: {output_path}")

