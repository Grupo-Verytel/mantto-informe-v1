"""
Generador Sección 3: Informes de Medición de Niveles de Servicio (ANS)
Tipo: 🟨 GENERACIÓN PROGRAMÁTICA (python-docx)

SECCIÓN CRÍTICA - Impacto contractual y financiero
- Umbral contractual: 98.9% de disponibilidad
- Si no cumple: calcular penalidades y explicar causas

Subsecciones:
- 3.1 Penalidad de ANS
- 3.2 Consolidado ANS (histórico y gráficos)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
import json
import calendar
from .base import GeneradorSeccion
import config


class GeneradorSeccion3(GeneradorSeccion):
    """Genera la Sección 3: Informes de Medición de Niveles de Servicio (ANS)"""
    
    # Umbrales contractuales
    UMBRAL_ANS = 98.9           # Porcentaje mínimo requerido
    UMBRAL_AMARILLO = 97.9      # Zona de alerta
    
    # Colores corporativos
    COLOR_AZUL_OSCURO = RGBColor(31, 78, 121)   # Encabezados principales
    COLOR_AZUL_MEDIO = RGBColor(46, 117, 182)   # Subsecciones
    COLOR_GRIS = RGBColor(64, 64, 64)           # Subtítulos
    COLOR_VERDE = RGBColor(0, 176, 80)          # ANS cumplido
    COLOR_ROJO = RGBColor(192, 0, 0)            # ANS no cumplido
    COLOR_AMARILLO = RGBColor(255, 192, 0)      # Zona de alerta
    
    # Colores de fondo para semáforo
    COLOR_VERDE_CLARO = RGBColor(198, 239, 206)
    COLOR_AMARILLO_CLARO = RGBColor(255, 242, 204)
    COLOR_ROJO_CLARO = RGBColor(255, 199, 206)
    
    @property
    def nombre_seccion(self) -> str:
        return "3. INFORMES DE MEDICIÓN DE NIVELES DE SERVICIO (ANS)"
    
    @property
    def template_file(self) -> str:
        return "seccion_3_ans.docx"  # No se usa, pero debe existir para compatibilidad
    
    def __init__(self, anio: int, mes: int):
        super().__init__(anio, mes)
        self.datos: Dict[str, Any] = {}
        self.doc: Optional[Document] = None
        self.disponibilidad: float = 0.0
        self.cumple_ans: bool = False
    
    def _configurar_estilos(self):
        """Configura los estilos del documento"""
        if self.doc is None:
            return
        
        # Configurar márgenes (1.27 cm = 0.5 pulgadas)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
        
        # Estilo normal
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Título Sección (3.)
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = self.COLOR_AZUL_OSCURO
        
        # Subsecciones (3.1, 3.2)
        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.font.color.rgb = self.COLOR_AZUL_MEDIO
        
        # Subtítulos
        h3 = self.doc.styles['Heading 3']
        h3.font.name = 'Arial'
        h3.font.size = Pt(11)
        h3.font.bold = True
        h3.font.color.rgb = self.COLOR_GRIS
    
    def _aplicar_sombreado_celda(self, cell, color: RGBColor):
        """Aplica color de fondo a celda de tabla"""
        shading_elm = OxmlElement('w:shd')
        # RGBColor es una tupla, acceder con índices
        r = color[0]
        g = color[1]
        b = color[2]
        
        hex_color = f'{r:02X}{g:02X}{b:02X}'
        shading_elm.set(qn('w:fill'), hex_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def _centrar_celda_vertical(self, cell):
        """Centra verticalmente el contenido de una celda"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
    
    def _aplicar_color_disponibilidad(self, cell, porcentaje: float):
        """Aplica color según semáforo de disponibilidad"""
        if porcentaje >= self.UMBRAL_ANS:
            # Verde: cumple
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = self.COLOR_VERDE
            self._aplicar_sombreado_celda(cell, self.COLOR_VERDE_CLARO)
        elif porcentaje >= self.UMBRAL_AMARILLO:
            # Amarillo: zona de alerta
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = self.COLOR_AMARILLO
            self._aplicar_sombreado_celda(cell, self.COLOR_AMARILLO_CLARO)
        else:
            # Rojo: crítico
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = self.COLOR_ROJO
            self._aplicar_sombreado_celda(cell, self.COLOR_ROJO_CLARO)
    
    def _agregar_parrafo(self, texto: str, justificado: bool = True, 
                        negrita: bool = False, color: RGBColor = None, tamano: int = 11):
        """Agrega un párrafo de texto"""
        p = self.doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = negrita
        run.font.name = 'Arial'
        run.font.size = Pt(tamano)
        if color:
            run.font.color.rgb = color
        if justificado:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        return p
    
    def _calcular_penalidad(self) -> Dict[str, Any]:
        """
        Calcula la penalidad por incumplimiento de ANS
        Escala contractual:
        - Déficit ≤ 0.5%: 0.5% del contrato
        - Déficit ≤ 1.0%: 1.0% del contrato
        - Déficit ≤ 1.5%: 1.5% del contrato
        - Déficit > 1.5%: 2.0% del contrato
        """
        if self.disponibilidad >= self.UMBRAL_ANS:
            return {
                'aplica': False,
                'deficit': 0.0,
                'porcentaje_penalidad': 0.0,
                'valor_penalidad': 0.0
            }
        
        deficit = self.UMBRAL_ANS - self.disponibilidad
        valor_mensual = self.datos.get('valor_mensual_contrato', 500000000)
        
        # Escala de penalidad según especificaciones
        if deficit <= 0.5:
            porcentaje_penalidad = 0.5
        elif deficit <= 1.0:
            porcentaje_penalidad = 1.0
        elif deficit <= 1.5:
            porcentaje_penalidad = 1.5
        else:
            porcentaje_penalidad = 2.0
        
        valor_penalidad = valor_mensual * (porcentaje_penalidad / 100)
        
        return {
            'aplica': True,
            'deficit': deficit,
            'porcentaje_penalidad': porcentaje_penalidad,
            'valor_penalidad': valor_penalidad
        }
    
    def _agregar_titulo_seccion(self):
        """Título principal: 3. INFORMES DE MEDICIÓN DE NIVELES DE SERVICIO (ANS)"""
        self.doc.add_heading("3. INFORMES DE MEDICIÓN DE NIVELES DE SERVICIO (ANS)", level=1)
    
    def _agregar_introduccion(self):
        """Introducción con fórmula de disponibilidad"""
        mes = self.datos.get('mes', config.MESES[self.mes])
        anio = self.datos.get('anio', self.anio)
        
        texto = (
            f"El contrato establece un Acuerdo de Nivel de Servicio (ANS) mínimo del {self.UMBRAL_ANS}% "
            f"de disponibilidad mensual para el sistema de videovigilancia. La disponibilidad "
            f"se calcula mediante la siguiente fórmula:\n\n"
            f"Disponibilidad (%) = (Horas Operativas / Horas Totales del Mes) × 100\n\n"
            f"A continuación se presenta el análisis de cumplimiento del ANS para el mes "
            f"de {mes} de {anio}."
        )
        self._agregar_parrafo(texto)
        
        # Fórmula destacada
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Disponibilidad (%) = (Horas Operativas / Horas Totales del Mes) × 100")
        run.italic = True
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = self.COLOR_GRIS
        self.doc.add_paragraph()
    
    def _agregar_3_1_penalidad_ans(self):
        """Subsección 3.1 completa"""
        self.doc.add_heading("3.1. PENALIDAD DE ANS", level=2)
        
        # Datos del periodo
        self._agregar_datos_periodo()
        
        # Resultado ANS
        self._agregar_resultado_ans()
        
        # Análisis de cumplimiento
        self._agregar_analisis_cumplimiento()
        
        # Tabla de localidades
        self._agregar_tabla_localidades()
        
        # Cálculo de penalidad (solo si no cumple)
        if not self.cumple_ans:
            self._agregar_calculo_penalidad()
    
    def _agregar_datos_periodo(self):
        """Tabla resumen del periodo"""
        self._agregar_parrafo("Datos del Periodo:", negrita=True)
        
        total_camaras = self.datos.get('total_camaras', 0)
        dias_mes = self.datos.get('dias_mes', calendar.monthrange(self.anio, self.mes)[1])
        horas_totales = self.datos.get('horas_totales', total_camaras * dias_mes * 24)
        horas_operativas = self.datos.get('horas_operativas', 0)
        horas_no_operativas = self.datos.get('horas_no_operativas', 0)
        
        # Crear tabla
        tabla = self.doc.add_table(rows=1, cols=2)
        tabla.style = 'Table Grid'
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = "Concepto"
        hdr_cells[1].text = "Valor"
        
        for i in range(2):
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            self._aplicar_sombreado_celda(hdr_cells[i], self.COLOR_GRIS)
            self._centrar_celda_vertical(hdr_cells[i])
        
        # Filas de datos
        filas = [
            ["Total de cámaras en operación", f"{total_camaras:,}"],
            ["Horas totales del periodo", f"{horas_totales:,} hrs"],
            ["Horas operativas", f"{horas_operativas:,} hrs"],
            ["Horas no operativas", f"{horas_no_operativas:,} hrs"],
            ["Disponibilidad calculada", f"{self.disponibilidad:.2f}%"]
        ]
        
        for fila in filas:
            row_cells = tabla.add_row().cells
            row_cells[0].text = fila[0]
            row_cells[1].text = fila[1]
            
            for i in range(2):
                for paragraph in row_cells[i].paragraphs:
                    alineacion = WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.alignment = alineacion
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                self._centrar_celda_vertical(row_cells[i])
            
            # Aplicar color a la última fila según cumplimiento
            if fila[0] == "Disponibilidad calculada":
                self._aplicar_color_disponibilidad(row_cells[1], self.disponibilidad)
        
        # Ajustar anchos
        for cell in tabla.columns[0].cells:
            cell.width = Cm(8.0)
        for cell in tabla.columns[1].cells:
            cell.width = Cm(6.0)
        
        self.doc.add_paragraph()
    
    def _agregar_resultado_ans(self):
        """Resultado: ✓ CUMPLE o ✗ NO CUMPLE"""
        self._agregar_parrafo("Resultado del ANS:", negrita=True)
        
        if self.cumple_ans:
            texto = f"✓ ANS CUMPLIDO - Disponibilidad: {self.disponibilidad:.2f}% (Umbral: {self.UMBRAL_ANS}%)"
            self._agregar_parrafo(texto, negrita=True, color=self.COLOR_VERDE, tamano=11)
        else:
            texto = f"✗ ANS NO CUMPLIDO - Disponibilidad: {self.disponibilidad:.2f}% (Umbral: {self.UMBRAL_ANS}%)"
            self._agregar_parrafo(texto, negrita=True, color=self.COLOR_ROJO, tamano=11)
        
        self.doc.add_paragraph()
    
    def _agregar_analisis_cumplimiento(self):
        """Análisis cualitativo (condicional)"""
        self._agregar_parrafo("Análisis de Cumplimiento:", negrita=True)
        
        mes = self.datos.get('mes', config.MESES[self.mes])
        anio = self.datos.get('anio', self.anio)
        
        if self.cumple_ans:
            factores = self.datos.get('factores_cumplimiento', [
                "mantenimiento preventivo programado",
                "respuesta rápida ante incidencias",
                "bajo índice de fallas de conectividad"
            ])
            factores_texto = "\n".join([f"• {f}" for f in factores])
            
            texto = (
                f"Durante el mes de {mes}, el sistema de videovigilancia alcanzó una disponibilidad "
                f"del {self.disponibilidad:.2f}%, superando el umbral contractual del {self.UMBRAL_ANS}%. "
                f"Este resultado se logró gracias a las actividades preventivas de mantenimiento, "
                f"la respuesta oportuna ante fallas, y la gestión eficiente de los recursos técnicos. "
                f"Los principales factores que contribuyeron al cumplimiento fueron:\n\n{factores_texto}"
            )
        else:
            causas = self.datos.get('causas_incumplimiento', [
                "falla masiva de conectividad en zona sur (15 horas)",
                "corte de energía prolongado en Kennedy (8 horas)"
            ])
            acciones = self.datos.get('acciones_correctivas', [
                "instalación de enlaces de respaldo",
                "implementación de UPS adicionales"
            ])
            
            deficit = self.UMBRAL_ANS - self.disponibilidad
            causas_texto = "\n".join([f"• {c}" for c in causas])
            acciones_texto = "\n".join([f"• {a}" for a in acciones])
            
            texto = (
                f"Durante el mes de {mes}, el sistema de videovigilancia registró una disponibilidad "
                f"del {self.disponibilidad:.2f}%, presentando un déficit de {deficit:.2f}% respecto al umbral "
                f"contractual. Las principales causas del incumplimiento fueron:\n\n{causas_texto}\n\n"
                f"Se han implementado las siguientes acciones correctivas:\n\n{acciones_texto}"
            )
        
        self._agregar_parrafo(texto)
        self.doc.add_paragraph()
    
    def _agregar_tabla_localidades(self):
        """Tabla con semáforo por localidad"""
        self._agregar_parrafo("Disponibilidad por Localidad:", negrita=True)
        
        localidades = self.datos.get('disponibilidad_por_localidad', [])
        
        if not localidades:
            self._agregar_parrafo("No se registraron datos de disponibilidad por localidad durante este periodo.")
            return
        
        # Crear tabla
        tabla = self.doc.add_table(rows=1, cols=5)
        tabla.style = 'Table Grid'
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados
        encabezados = ["Localidad", "Cámaras", "Hrs Operativas", "Hrs No Operativas", "Disponibilidad (%)"]
        hdr_cells = tabla.rows[0].cells
        
        for i, texto in enumerate(encabezados):
            hdr_cells[i].text = texto
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            self._aplicar_sombreado_celda(hdr_cells[i], self.COLOR_AZUL_OSCURO)
            self._centrar_celda_vertical(hdr_cells[i])
        
        # Filas de datos
        for loc in localidades:
            row_cells = tabla.add_row().cells
            disponibilidad = loc.get('disponibilidad', loc.get('disponibilidad_porcentaje', 0))
            
            fila = [
                loc.get('nombre', loc.get('localidad', '')),
                str(loc.get('camaras', loc.get('total_camaras', 0))),
                f"{loc.get('horas_operativas', 0):,}",
                f"{loc.get('horas_no_operativas', 0):,}",
                f"{disponibilidad:.2f}%"
            ]
            
            for i, texto in enumerate(fila):
                row_cells[i].text = texto
                for paragraph in row_cells[i].paragraphs:
                    alineacion = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.alignment = alineacion
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                self._centrar_celda_vertical(row_cells[i])
            
            # Aplicar semáforo a la columna de disponibilidad
            self._aplicar_color_disponibilidad(row_cells[4], disponibilidad)
        
        # Ajustar anchos
        anchos = [Cm(4.5), Cm(2.5), Cm(3.0), Cm(3.5), Cm(3.5)]
        for i, ancho in enumerate(anchos):
            for cell in tabla.columns[i].cells:
                cell.width = ancho
        
        self.doc.add_paragraph()
    
    def _agregar_calculo_penalidad(self):
        """Muestra cálculo de penalidad (solo si no cumple)"""
        self._agregar_parrafo("Cálculo de Penalidad:", negrita=True)
        
        penalidad = self._calcular_penalidad()
        
        if not penalidad['aplica']:
            self._agregar_parrafo(
                "No aplica penalidad. El ANS fue cumplido satisfactoriamente.",
                color=self.COLOR_VERDE
            )
            return
        
        texto = (
            f"Por el incumplimiento del ANS contractual, se calcula una penalidad de "
            f"{penalidad['porcentaje_penalidad']:.1f}% sobre el valor mensual del contrato, "
            f"equivalente a ${penalidad['valor_penalidad']:,.0f} pesos colombianos."
        )
        self._agregar_parrafo(texto, color=self.COLOR_ROJO)
        
        # Tabla de penalidad
        tabla = self.doc.add_table(rows=1, cols=2)
        tabla.style = 'Table Grid'
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = "Concepto"
        hdr_cells[1].text = "Valor"
        
        for i in range(2):
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            self._aplicar_sombreado_celda(hdr_cells[i], self.COLOR_ROJO)
            self._centrar_celda_vertical(hdr_cells[i])
        
        # Filas de datos
        filas = [
            ["Déficit de disponibilidad", f"{penalidad['deficit']:.2f}%"],
            ["Porcentaje de penalidad", f"{penalidad['porcentaje_penalidad']:.1f}%"],
            ["Valor de la penalidad", f"${penalidad['valor_penalidad']:,.0f}"]
        ]
        
        for fila in filas:
            row_cells = tabla.add_row().cells
            row_cells[0].text = fila[0]
            row_cells[1].text = fila[1]
            
            for i in range(2):
                for paragraph in row_cells[i].paragraphs:
                    alineacion = WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.alignment = alineacion
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                self._centrar_celda_vertical(row_cells[i])
            
            # Aplicar color rojo claro a todas las filas
            for cell in row_cells:
                self._aplicar_sombreado_celda(cell, self.COLOR_ROJO_CLARO)
        
        # Ajustar anchos
        for cell in tabla.columns[0].cells:
            cell.width = Cm(8.0)
        for cell in tabla.columns[1].cells:
            cell.width = Cm(6.0)
        
        self.doc.add_paragraph()
    
    def _agregar_3_2_consolidado_ans(self):
        """Subsección 3.2 completa"""
        self.doc.add_heading("3.2. CONSOLIDADO ANS", level=2)
        
        self._agregar_parrafo(
            "A continuación se presenta el histórico de cumplimiento del ANS desde el inicio del contrato:"
        )
        self.doc.add_paragraph()
        
        # Tabla histórico
        self._agregar_tabla_historico()
        
        # Resumen acumulado
        self._agregar_resumen_acumulado()
        
        # Nota sobre gráficos
        self._agregar_parrafo(
            "Nota: Los gráficos de tendencia del ANS se encuentran en el Anexo 3.1 del presente informe.",
            negrita=False
        )
    
    def _agregar_tabla_historico(self):
        """Tabla histórico mensual"""
        historico = self.datos.get('historico_ans', [])
        
        if not historico:
            self._agregar_parrafo("No se cuenta con histórico de ANS para este periodo.")
            return
        
        # Crear tabla
        tabla = self.doc.add_table(rows=1, cols=5)
        tabla.style = 'Table Grid'
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados
        encabezados = ["Mes", "Disponibilidad (%)", "Umbral (%)", "Estado", "Observaciones"]
        hdr_cells = tabla.rows[0].cells
        
        for i, texto in enumerate(encabezados):
            hdr_cells[i].text = texto
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            self._aplicar_sombreado_celda(hdr_cells[i], self.COLOR_AZUL_MEDIO)
            self._centrar_celda_vertical(hdr_cells[i])
        
        # Filas de datos
        for h in historico:
            row_cells = tabla.add_row().cells
            disponibilidad = h.get('disponibilidad', 0)
            cumple = disponibilidad >= self.UMBRAL_ANS
            
            fila = [
                h.get('mes', ''),
                f"{disponibilidad:.2f}%",
                f"{self.UMBRAL_ANS}%",
                "✓ Cumple" if cumple else "✗ No Cumple",
                h.get('observaciones', '-')
            ]
            
            for i, texto in enumerate(fila):
                row_cells[i].text = texto
                for paragraph in row_cells[i].paragraphs:
                    alineacion = WD_ALIGN_PARAGRAPH.CENTER if i < 4 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.alignment = alineacion
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                self._centrar_celda_vertical(row_cells[i])
            
            # Aplicar color a la columna de estado
            if cumple:
                for paragraph in row_cells[3].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = self.COLOR_VERDE
                self._aplicar_sombreado_celda(row_cells[3], self.COLOR_VERDE_CLARO)
            else:
                for paragraph in row_cells[3].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = self.COLOR_ROJO
                self._aplicar_sombreado_celda(row_cells[3], self.COLOR_ROJO_CLARO)
        
        # Ajustar anchos
        anchos = [Cm(3.0), Cm(3.0), Cm(2.5), Cm(2.5), Cm(6.0)]
        for i, ancho in enumerate(anchos):
            for cell in tabla.columns[i].cells:
                cell.width = ancho
        
        self.doc.add_paragraph()
    
    def _agregar_resumen_acumulado(self):
        """Resumen estadístico acumulado"""
        self._agregar_parrafo("Resumen Acumulado:", negrita=True)
        
        historico = self.datos.get('historico_ans', [])
        
        if not historico:
            self._agregar_parrafo("No se cuenta con datos suficientes para el resumen acumulado.")
            return
        
        meses_cumplidos = sum(1 for h in historico if h.get('disponibilidad', 0) >= self.UMBRAL_ANS)
        total_meses = len(historico)
        promedio_disponibilidad = sum(h.get('disponibilidad', 0) for h in historico) / total_meses if total_meses > 0 else 0
        
        # Crear tabla
        tabla = self.doc.add_table(rows=1, cols=2)
        tabla.style = 'Table Grid'
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = "Indicador"
        hdr_cells[1].text = "Valor"
        
        for i in range(2):
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            self._aplicar_sombreado_celda(hdr_cells[i], self.COLOR_GRIS)
            self._centrar_celda_vertical(hdr_cells[i])
        
        # Filas de datos
        filas = [
            ["Total meses evaluados", str(total_meses)],
            ["Meses con ANS cumplido", str(meses_cumplidos)],
            ["Meses con ANS no cumplido", str(total_meses - meses_cumplidos)],
            ["Porcentaje de cumplimiento", f"{(meses_cumplidos/total_meses*100):.1f}%" if total_meses > 0 else "N/A"],
            ["Disponibilidad promedio", f"{promedio_disponibilidad:.2f}%"]
        ]
        
        for fila in filas:
            row_cells = tabla.add_row().cells
            row_cells[0].text = fila[0]
            row_cells[1].text = fila[1]
            
            for i in range(2):
                for paragraph in row_cells[i].paragraphs:
                    alineacion = WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.alignment = alineacion
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                self._centrar_celda_vertical(row_cells[i])
        
        # Ajustar anchos
        for cell in tabla.columns[0].cells:
            cell.width = Cm(8.0)
        for cell in tabla.columns[1].cells:
            cell.width = Cm(6.0)
        
        self.doc.add_paragraph()
    
    def cargar_datos(self) -> None:
        """Carga los datos específicos de la sección 3 desde JSON"""
        archivo = config.FUENTES_DIR / f"ans_{self.mes}_{self.anio}.json"
        
        if archivo.exists():
            try:
                with open(archivo, 'r', encoding='utf-8') as f:
                    self.datos = json.load(f)
            except Exception as e:
                print(f"[WARNING] Error al cargar datos desde {archivo}: {e}")
                self.datos = self._datos_ejemplo()
        else:
            print(f"[WARNING] Archivo de datos no encontrado: {archivo}")
            self.datos = self._datos_ejemplo()
        
        # Calcular disponibilidad si no está en los datos
        if 'disponibilidad_porcentaje' not in self.datos:
            horas_totales = self.datos.get('horas_totales', 0)
            horas_operativas = self.datos.get('horas_operativas', 0)
            if horas_totales > 0:
                self.datos['disponibilidad_porcentaje'] = (horas_operativas / horas_totales) * 100
        
        # Calcular disponibilidad y cumplimiento
        self.disponibilidad = self.datos.get('disponibilidad_porcentaje', 0)
        self.cumple_ans = self.disponibilidad >= self.UMBRAL_ANS
        
        # Agregar mes y año a los datos para compatibilidad
        if 'mes' not in self.datos:
            self.datos['mes'] = config.MESES[self.mes]
        if 'anio' not in self.datos:
            self.datos['anio'] = self.anio
    
    def _datos_ejemplo(self) -> Dict[str, Any]:
        """Retorna datos de ejemplo para desarrollo"""
        dias_mes = calendar.monthrange(self.anio, self.mes)[1]
        total_camaras = 5824
        horas_totales = total_camaras * dias_mes * 24
        horas_operativas = int(horas_totales * 0.9917)  # 99.17%
        horas_no_operativas = horas_totales - horas_operativas
        
        return {
            "mes": config.MESES[self.mes],
            "anio": self.anio,
            "total_camaras": total_camaras,
            "dias_mes": dias_mes,
            "horas_totales": horas_totales,
            "horas_operativas": horas_operativas,
            "horas_no_operativas": horas_no_operativas,
            "disponibilidad_porcentaje": 99.17,
            "valor_mensual_contrato": 500000000,
            "disponibilidad_por_localidad": [],
            "historico_ans": [],
            "factores_cumplimiento": [
                "mantenimiento preventivo programado",
                "respuesta rápida ante incidencias",
                "bajo índice de fallas de conectividad"
            ]
        }
    
    def procesar(self) -> Dict[str, Any]:
        """Procesa los datos y retorna el contexto (no se usa en generación programática)"""
        return {}
    
    def generar(self) -> Document:
        """
        Genera el documento completo de la Sección 3
        Sobrescribe el método de la clase base para usar python-docx directamente
        """
        # Cargar datos si no se han cargado
        if not self.datos:
            self.cargar_datos()
        
        # Crear documento
        self.doc = Document()
        self._configurar_estilos()
        
        # Generar contenido
        self._agregar_titulo_seccion()
        self._agregar_introduccion()
        self._agregar_3_1_penalidad_ans()
        self._agregar_3_2_consolidado_ans()
        
        # Separador fin de sección
        self.doc.add_paragraph()
        p = self.doc.add_paragraph("═" * 60)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = self.doc.add_paragraph("Fin Sección 3 - Informes de Medición de Niveles de Servicio (ANS)")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        return self.doc
    
    def guardar(self, output_path: Path) -> None:
        """
        Genera y guarda la sección
        Sobrescribe el método de la clase base para usar python-docx
        """
        if self.doc is None:
            self.generar()
        
        self.doc.save(str(output_path))
        print(f"[OK] {self.nombre_seccion} guardada en: {output_path}")
