"""
Utilidades para crear tablas en documentos Word
"""
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_tabla_desde_dict(doc: Document, datos: List[Dict[str, Any]], encabezados: List[str]) -> None:
    """
    Crea una tabla en un documento Word a partir de una lista de diccionarios
    
    Args:
        doc: Documento Word
        datos: Lista de diccionarios con los datos
        encabezados: Lista de nombres de columnas
    """
    if not datos:
        return
    
    # Crear tabla
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = 'Light Grid Accent 1'
    
    # Agregar encabezados
    header_cells = tabla.rows[0].cells
    for i, encabezado in enumerate(encabezados):
        header_cells[i].text = encabezado
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Agregar datos
    for fila_datos in datos:
        row_cells = tabla.add_row().cells
        for i, encabezado in enumerate(encabezados):
            valor = fila_datos.get(encabezado, "")
            row_cells[i].text = str(valor) if valor is not None else ""
    
    return tabla

def crear_tabla_desde_lista(doc: Document, datos: List[List[Any]], encabezados: List[str] = None) -> None:
    """
    Crea una tabla en un documento Word a partir de una lista de listas
    
    Args:
        doc: Documento Word
        datos: Lista de listas con los datos (cada lista es una fila)
        encabezados: Lista opcional de nombres de columnas
    """
    if not datos:
        return
    
    num_cols = len(datos[0]) if datos else len(encabezados) if encabezados else 0
    if num_cols == 0:
        return
    
    # Crear tabla
    num_rows = len(datos) + (1 if encabezados else 0)
    tabla = doc.add_table(rows=num_rows, cols=num_cols)
    tabla.style = 'Light Grid Accent 1'
    
    # Agregar encabezados si existen
    if encabezados:
        header_cells = tabla.rows[0].cells
        for i, encabezado in enumerate(encabezados):
            header_cells[i].text = str(encabezado)
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        inicio_datos = 1
    else:
        inicio_datos = 0
    
    # Agregar datos
    for idx, fila_datos in enumerate(datos):
        row_cells = tabla.rows[inicio_datos + idx].cells
        for i, valor in enumerate(fila_datos):
            row_cells[i].text = str(valor) if valor is not None else ""
    
    return tabla


